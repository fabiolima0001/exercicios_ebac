 Código Python para gerar o documento
python
Copiar
Editar
from docx import Document

doc = Document()
doc.add_heading('Exercício - Técnicas de Testes', 0)
doc.add_heading('Funcionalidade: Cadastro de novos produtos - Loja EBAC', level=1)

# Exercício 1
doc.add_heading('Exercício 1: Particionamento de Equivalência', level=2)
doc.add_heading('RN01', level=3)
table = doc.add_table(rows=1, cols=2)
table.rows[0].cells[0].text = "Entrada"
table.rows[0].cells[1].text = "Saída"
for e, s in [("R$18,99", "Inválida"), ("R$19,00", "Válida"), ("R$32,00", "Válida"), ("R$99,00", "Válida"), ("R$99,01", "Inválida")]:
    row = table.add_row().cells
    row[0].text, row[1].text = e, s

doc.add_heading('RN02', level=3)
table = doc.add_table(rows=1, cols=3)
table.rows[0].cells[:] = ["Produto igual", "Data último cadastro", "Saída"]
for e in [("Sim", "40 dias atrás", "Válida (renovar)"),
          ("Sim", "10 dias atrás", "Inválida (não renovar)"),
          ("Não", "—", "Válida")]:
    row = table.add_row().cells
    row[0].text, row[1].text, row[2].text = e

doc.add_heading('RN03', level=3)
table = doc.add_table(rows=1, cols=2)
table.rows[0].cells[:] = ["Quantidade", "Saída"]
for q, s in [("50", "Válida"), ("100", "Válida"), ("101", "Inválida")]:
    row = table.add_row().cells
    row[0].text, row[1].text = q, s

# Exercício 2
doc.add_heading('Exercício 2: Valor Limite', level=2)
doc.add_heading('RN01', level=3)
table = doc.add_table(rows=1, cols=2)
table.rows[0].cells[:] = ["Entrada", "Saída"]
for e, s in [("R$18,99", "Inválida"), ("R$19,00", "Válida"), ("R$19,01", "Válida"),
             ("R$98,99", "Válida"), ("R$99,00", "Válida"), ("R$99,01", "Inválida")]:
    row = table.add_row().cells
    row[0].text, row[1].text = e, s

doc.add_heading('RN03', level=3)
table = doc.add_table(rows=1, cols=2)
table.rows[0].cells[:] = ["Quantidade", "Saída"]
for q, s in [("99", "Válida"), ("100", "Válida"), ("101", "Inválida")]:
    row = table.add_row().cells
    row[0].text, row[1].text = q, s

# Exercício 3
doc.add_heading('Exercício 3: Tabela de Decisão', level=2)
doc.add_paragraph("Regras: RN02 - Renovar produtos iguais cadastrados há mais de 30 dias\nRN03 - Permitir no máximo 100 itens por vez")
table = doc.add_table(rows=1, cols=5)
table.rows[0].cells[:] = ["Produto Igual", "Dias desde último cadastro", "Qtde Produtos ≤ 100", "Ação: Renovar", "Ação: Cadastrar"]
linhas = [
    ("Sim", "> 30 dias", "Sim", "Sim", "Sim"),
    ("Sim", "≤ 30 dias", "Sim", "Não", "Sim"),
    ("Sim", "> 30 dias", "Não", "Sim", "Não"),
    ("Não", "—", "Sim", "Não", "Sim"),
    ("Não", "—", "Não", "Não", "Não")
]
for l in linhas:
    row = table.add_row().cells
    for i, val in enumerate(l):
        row[i].text = val

doc.add_heading("Observações Finais", level=2)
doc.add_paragraph(
    "Cada técnica cobre diferentes aspectos da validação de requisitos. "
    "As regras foram abordadas com testes mínimos para 100% de cobertura. "
    "Recomenda-se automação desses testes para garantir consistência."
)
